"""
The flask application for the career-flow application
"""
# importing required python libraries
from flask import Flask, jsonify, request, send_file
from flask_mongoengine import MongoEngine
from flask_cors import CORS, cross_origin
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from webdriver_manager.chrome import ChromeDriverManager
import pandas as pd
from datetime import datetime, timedelta
import yaml
import hashlib
import uuid
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from flask_pymongo import PyMongo, ObjectId
from io import BytesIO
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from apscheduler.schedulers.background import BackgroundScheduler
from datetime import datetime, timedelta
from collections import defaultdict
import sys

import json
import bcrypt
from flask_pymongo import PyMongo
from flask import Flask, request, jsonify
from flask_pymongo import PyMongo
from pymongo import MongoClient
import mongomock
from datetime import datetime, timedelta, timezone
from flask_jwt_extended import create_access_token,get_jwt,get_jwt_identity, \
                               unset_jwt_cookies, jwt_required, JWTManager
from datetime import datetime, timedelta
from functools import reduce
from bson import json_util 
from pymongo import MongoClient
from flasgger import Swagger
from flask_mail import Mail, Message
import smtplib
from smtplib import SMTPAuthenticationError

from io import BytesIO


existing_endpoints = ["/applications", "/resume", "/dashboard", "/contacts", "/token", "/register"]




def create_app():
    """
    Creates a server hosted on localhost

    :return: Flask object
    """
    app = Flask(__name__)
    app.secret_key = 'secret'
    app.config["JWT_SECRET_KEY"] = "softwareEngineering"
    app.config["JWT_ACCESS_TOKEN_EXPIRES"] = timedelta(hours=1)
    # make flask support CORS
    CORS(app)
    app.config['MAIL_SERVER'] = 'smtp.mail.yahoo.com'
    app.config['MAIL_PORT'] = 587
    app.config['MAIL_USERNAME'] = 'contactus_burnout@yahoo.com'  # Your Gmail email address
    app.config['MAIL_PASSWORD'] = 'fpfkvuxxubnmzahw'  # Your Gmail password or app-specific password
    app.config['MAIL_USE_TLS'] = True
    app.config['MAIL_USE_SSL'] = False
        
    return app

app = create_app()
# with open("application.yml") as f:
#     info = yaml.load(f, Loader=yaml.FullLoader)
#     username = info["username"]
#     password = info["password"]
#     app.config["MONGODB_SETTINGS"] = {
#         "db": "appTracker",
#         # "host": os.getenv("db_username"),
#         "host": "localhost",
#     }

app.config["MONGODB_SETTINGS"] = {
        "db": "appTracker",
        # "host": os.getenv("db_username"),
        "host": "localhost",
    }
db = MongoEngine()
db.init_app(app)
mail = Mail(app)
jwt = JWTManager(app)

@app.route("/")
@cross_origin()
def health_check():
    """
    Return 200. Just a health check method to check if flask application is running.
    Corresponds to API route /
    """
    return jsonify({"message": "Server up and running"}), 200
     

@app.route('/token', methods=["POST"])
def create_token():
    """
    Login function that creates a jwt_token for the user and returns it
    Corresponds to API route /token
    """
    print('This is standard output')
    try:
        email = request.json.get("email", None)
        password = request.json.get("password", None)
        print(password)
        user = Users.objects(
            email=email, password=password
        ).first()
        if user is not None:
            access_token = create_access_token(identity=email)
            return jsonify({"message": "Login successful", "access_token":access_token})
        else:
            print("Invalid email or password")
            return jsonify({"message": "Invalid email or password"}),401
    except Exception as err:
        print( err)
        return jsonify({"error": "Internal server error"}), 500
    
@app.route("/register", methods=["POST"])
def register():
    """
    Signup function that creates a new user object and adds it to the database.
    Corresponds to API route /register
    """
    email = request.json.get('email', None)
    password = request.json.get('password', None)
    first_name = request.json.get('firstName', None)
    last_name = request.json.get('lastName', None)
    print(email)
    new_document = {
    "email": email,
    "password": password,
    "first_name": first_name,
    "last_name": last_name,
    }
    query = {
        "email": email,
    }
    try:
        user = Users(
            firstName=first_name,
            lastName=last_name,
            email=email,
            password=password,
            applications=[]
        )
        user.save()
        return jsonify(user.to_json()), 200
    except Exception as e:
        response = jsonify({"msg": "register failed"})
    return response
    
@app.route("/logout", methods=["POST"])
def logout():
    """
    Logout the user and clear their session.
    Corresponds to API route /logout
    tags:
    - User Logout
    responses:
    200:
        description: Logout successful
    """
    response = jsonify({"msg": "logout successful"})
    unset_jwt_cookies(response)
    return response

# get data from the CSV file for rendering root page
@app.route("/applications", methods=["GET"])
@jwt_required()
def get_data():
    """
    Gets user's applications from the database.
    Corresponds to API route GET /application
    
    :return: JSON object with application data as a list of application objects
    """
    try:
        user_email = get_jwt_identity()
        users = Users.objects()
        user = users.filter(email=user_email).first()
        applications = user["applications"]
        return jsonify(applications)
    except Exception as e:
        print(e)
        return jsonify({"error": "Internal server error"}), 500

@app.route("/applications", methods=["POST"])
@jwt_required()
def add_application():
    """
    Adds a new job application for the user with the details from the request body.
    Corresponds to API route POST /applications
    
    :return: JSON object with status and message
    """
    
    try:
        current_user_email = get_jwt_identity()
        print("User: ------ ", current_user_email)
        users = Users.objects()
        user = users.filter(email=current_user_email).first()
        print("User: ------ ", user.email)
        print(request.json.get('date', None))
        current_application = {
            "id": get_new_application_id(user),
            "jobTitle": request.json.get('jobTitle', None),
            "companyName": request.json.get('companyName', None),
            "date": request.json.get('date', None)[:10],
            "jobLink": request.json.get('jobLink', None),
            "location": request.json.get('location', None),
            "stage": request.json.get("status", "1"),
                "notes": "",
        }
        applications = user["applications"] + [current_application]
        user["applications"] = applications
        user.save()
        return jsonify(current_application), 200
    except Exception as ex:
        return jsonify({"error": ex}), 500

@app.route("/applications/<int:application_id>", methods=["PUT"])
@jwt_required()
def update_application(application_id):
    """
    Updates the given job application for the user with the details from the request body
    Corresponds to API route PUT /applications/<application_id>
    
    :param application_id: Application id to be modified
    
    :return: JSON object with status and message
    """
    try:
        user_email = get_jwt_identity()
        request_data = request.json.get("application", {})
        print("HERE:")
        print(request_data)
        user = Users.objects(email=user_email).first()
        if not user:
            return jsonify({"error": "User not found"}), 404
        application_to_update = None
        for app in user.applications:
            if app["id"] == application_id:
                for key in request_data:
                    app[key] = request_data[key]
                application_to_update = app
                break
        if application_to_update is None:
            return jsonify({"error": "Application not found"}), 404
        user.save()
        return jsonify(application_to_update), 200
    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500
    
@app.route("/applications/<int:application_id>", methods=["DELETE"])
@jwt_required()
def delete_application(application_id):
    """
    Deletes the given job application for the user.
    Corresponds to API route DELETE /applications/<application_id>
    
    :param application_id: Application id to be deleted
    
    :return: JSON object with status and message
    """
    try:
        user_email = get_jwt_identity()
        user = Users.objects(email=user_email).first()
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        # Filter out the application with the given ID
        user.applications = [app for app in user.applications if app['id'] != application_id]
        user.save()
        # You can add an email notification here if needed
        return jsonify({"message": "Application deleted"}), 200
    except Exception as e:
        print(e)
        return jsonify({"error": "Internal server error"}), 500
        
@app.route("/applications/<application_id>/notes", methods=["PUT"])
@jwt_required()
def update_application_notes(application_id):
    try:
        current_user_email = get_jwt_identity()
        user = Users.objects(email=current_user_email).first()

        if not user:
            return jsonify({"error": "User not found"}), 404

        # Parse the application ID as an integer or appropriate format
        application_id = int(application_id)

        # Find the application in the user's applications list
        application_to_update = next((app for app in user.applications if app.get('id') == application_id), None)

        if not application_to_update:
            return jsonify({"error": "Application not found"}), 404

        # Update the note for the application
        note = request.json.get('notes')
        application_to_update['notes'] = note
        user.save()

        return jsonify(application_to_update), 200
    except Exception as ex:
        return jsonify({"error": str(ex)}), 500

@app.route("/resume", methods=["POST"])
@jwt_required()
def upload_resume():
    """
    Uploads resume file or updates an existing resume for the user
    Corresponds to API route POST /resume
    
    :return: JSON object with status and message
    """
    try:
        current_user = get_jwt_identity()
        user = Users.objects(email=current_user).first()
        try:
            file = request.files["file"]
        except:
            return jsonify({"error": "No resume file found in the input"}), 400

        if not hasattr(user, 'resume'):
            # There is no file
            user.resume.put(file.stream, content_type=file.content_type, filename=file.filename)
            user.save()
            return jsonify({"message": "resume successfully uploaded"}), 200
        else:
            # There is a file, we are replacing it
            user.resume.replace(file.stream, content_type=file.content_type, filename=file.filename)
            user.save()
            return jsonify({"message": "resume successfully replaced"}), 200
    except Exception as e:
        print(e)
        return jsonify({"error": "Internal server error"}), 500
    
@app.route('/sendEmail', methods=['POST'])
@jwt_required()
def send_email():
    """
    Send an email with the shared application details.
    Corresponds to API route POST /sendEmail
    """
    data = request.get_json()
    email = data.get('email')
    taskDetails = data.get('taskDetails')
    
    try:
        msg = Message("Shared Task Details", sender="contactus_burnout@yahoo.com", recipients=[email])
        msg.body = f"Task Details:\n{taskDetails}"
        mail.send(msg)
        return jsonify({"message": "Email sent successfully"}), 200
    except SMTPAuthenticationError as e:
        return jsonify({"error": "Authentication error. Check your credentials"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500 

@app.route("/downloadresume", methods=["GET"])
@jwt_required()
def get_resume():
    """
    Retrieves the resume file for the user
    Corresponds to API route GET /downloadresume
    
    :return: response with file
    """
    try:
        current_user = get_jwt_identity()
        user = Users.objects(email=current_user).first()
        try:
            if len(user.resume.read()) == 0:
                raise FileNotFoundError
            else:
                user.resume.seek(0)
        except:
            return jsonify({"error": "resume could not be found"}), 400
    
        file_like_object = BytesIO(user.resume.read())
        if hasattr(user, 'resume'):
            response = send_file(
            file_like_object,
            mimetype="application/pdf",
            download_name = user.resume.filename,
            as_attachment=True,
            )
            response.headers["x-filename"] = "resume.pdf"
            response.headers["Access-Control-Expose-Headers"] = "x-filename"
            return response, 200
        
    except Exception as ex:
            print(ex)
            return jsonify({"error": "Internal server error"}), 500

@app.route("/fetchresume", methods=["GET"])
@jwt_required()
def fetch_resume():
    """
    Retrieves the resume file for the user and sends it back in a way that
    can be displayed in an iframe on the client-side.
    Corresponds to API route GET /fetchresume
    
    :return: response with file
    """
    try:
        current_user = get_jwt_identity()
        user = Users.objects(email=current_user).first()
        if not hasattr(user, 'resume') or user.resume.length == 0:
            return jsonify({"error": "No resume uploaded"}), 404
        user.resume.seek(0)  # Reset the file pointer to the beginning
        file_like_object = BytesIO(user.resume.read())
        return send_file(
            file_like_object,
            mimetype="application/pdf",
            as_attachment=False  # Set to False for inline display
        )
    except Exception as ex:
        print(ex)
        return jsonify({"error": "Internal server error"}), 500

@app.route("/dashboard", methods=["GET"])
@jwt_required()
def get_dashboard_data():
    """
    Gets user's stats data from the database
    Corresponds to API route GET /dashboard
    
    :return: JSON object with stats data
    """
    try:
        user_email = get_jwt_identity()
        users = Users.objects()         
        user = users.filter(email=user_email).first()
        applications = user["applications"]
        print("Task 1/3: Getting Job Application Status Counts")
        job_app_status, applications_created, interviews_completed = get_job_app_status(applications)
        print("Task 2/3: Getting Job Applications Created in the last 6 months")
        six_months_job_count = get_last_six_months_job_counts(applications)
        print("Task 3/3: Getting latest 4 applications")
        last_four_apps = get_last_four_jobs(applications)
        contacts_saved = len(user["contacts"])
        notes_taken = len([1 for app in applications if "notes" in app and len(app["notes"]>0)])
        return jsonify({
            "six_months_jobs_count": six_months_job_count,
            "job_applications_status": job_app_status,
            "applications_created": applications_created,
            "interviews_completed": interviews_completed,
            "contacts_saved": contacts_saved,
            "notes_taken": notes_taken,
            "last_four_apps": last_four_apps
        }),200
    except Exception as e:
        print(e)
        return jsonify({"error": "Internal server error", "error msg":e}), 500
    
@app.route("/users/contacts", methods=["GET"])
@jwt_required()
def get_contacts():
    """
    Retrieves contacts for the logged-in user.
    Corresponds to API route GET /users/contacts
    
    :return: JSON object with user's contacts
    """
    try:
        current_user = get_jwt_identity()
        user = Users.objects(email=current_user).first()
        if not user:
            return jsonify({"error": "User not found"}), 404
        contacts = user.contacts
        return jsonify({"contacts": contacts}), 200
    except Exception as e:
        return jsonify({"error": "Unable to fetch contacts"}), 500

@app.route("/users/contacts", methods=["POST"])
@jwt_required()
def add_contact():
    """
    Adds a new contact for the user.
    Corresponds to API route POST /users/contacts
    
    :return: JSON object with status and message
    """
    try:
        current_user_id = get_jwt_identity()
        user = Users.objects(email=current_user_id).first()
        if not user:
            return jsonify({"error": "User not found"}), 404
        data = json.loads(request.data)
        new_contact = {
            "firstName": data["firstName"],
            "lastName": data["lastName"],
            "jobTitle": data.get("jobTitle", ""),
            "companyName": data.get("companyName", ""),
            "email": data.get("email", ""),
            "phone": data.get("phone", ""),
            "linkedin": data.get("linkedin", "")
        }
        user.contacts.append(new_contact)
        user.save()
        return jsonify({"message": "Contact added successfully", "contact": new_contact}), 200
    except Exception as e:
        print(e)
        return jsonify({"error": "Unable to add contact"}), 500


class Users(db.Document):
    """
    Users class. Holds full name, username, password, email, list of applications, list of contacts and a resume
    """

    firstName = db.StringField()
    lastName = db.StringField()
    email    = db.StringField()
    password = db.StringField()
    applications = db.ListField()
    contacts = db.ListField()
    resume = db.FileField()

class ResumeDocument(db.Document):
    """
    Resume document class. Stores the resume as a binary file object.
    """
    content = db.BinaryField()
    filename = db.StringField()

    def to_json(self):
        """
        Returns the user details in JSON object

        :return: JSON object
        """
        return { "fullName": self.firstName, "lastName": self.lastName,  "email": self.email}


def get_new_user_id():
    """
    Returns the next value to be used for new user

    :return: key with new user_id
    """
    user_objects = Users.objects()
    if len(user_objects) == 0:
        return 1

    new_id = 0
    for a in user_objects:
        new_id = max(new_id, a["id"])

    return new_id + 1


def get_new_application_id(user):
    """
    Returns the next value to be used for new application

    :param: user_id: User id of the active user

    :return: key with new application_id
    """

    if len(user["applications"]) == 0:
        return 1

    new_id = 0
    for a in user["applications"]:
        new_id = max(new_id, a["id"])

    return new_id + 1

def send_email(to_email, subject, message):
    # Set up your email and password here, or use environment variables
    gmail_user = "amoghmahesh14@gmail.com"
    gmail_password = os.getenv("email_password")

    msg = MIMEMultipart()
    msg['From'] = gmail_user
    msg['To'] = to_email
    msg['Subject'] = subject

    # Attach the message
    msg.attach(MIMEText(message, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(gmail_user, gmail_password)
        text = msg.as_string()
        server.sendmail(gmail_user, to_email, text)
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print("Email could not be sent. Error: {}".format(str(e)))

def email_reminders():
    users = Users.objects({})
    for user in users:
        current_applications =  user["applications"]
        for application in current_applications:
            if application["status"] != 3 or application["status"] != 4:
                try:
                    # Send an email reminder
                    to_email = user.email  # Use the user's email address
                    subject = "Job Application Reminder"
                    message = f"Hello {user.fullName},\n\nThe following job application has not been submitted yet.\nJob Title: {application['jobTitle']}\nCompany: {application['companyName']}\nApply By: {application['date']}\n\nBest regards,\nYour Application Tracker"
                    send_email(to_email, subject, message)
                except:
                    return jsonify({"error": "EMAIL wasn't sent"}), 400
                
sched = BackgroundScheduler(daemon=True)
sched.add_job(email_reminders,'interval',minutes=60)
sched.start()
def generate_pdf(data):
    doc = Document()

    # Set page margins to fit within one page
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(36)  # 0.5 inch
        section.right_margin = Pt(36)  # 0.5 inch
        section.top_margin = Pt(36)  # 0.5 inch
        section.bottom_margin = Pt(36)  # 0.5 inch

    # Helper function to add heading with format
    def add_heading_with_format(doc, text, font_size=16, is_bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if is_bold:
            run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.font.size = Pt(font_size)

    # Function to add details section
    def add_details_section(doc, section_title, details, is_bold_title=True):
        if section_title:
            add_heading_with_format(doc, section_title, font_size=14, is_bold=True)
        for detail in details:
            for key, value in detail.items():
                if key == "company":
                    p = doc.add_paragraph()
                    run = p.add_run(value)
                    run.bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif key == "project_title":
                    # Add the value of "project_title" with bold formatting
                    p = doc.add_paragraph()
                    run = p.add_run(value)
                    run.bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif key == "descriptionc":
                    # Add the value of "descriptionc" without "descriptionc" prefix
                    doc.add_paragraph(value, style="List Bullet")
                elif key != "descriptionc" and key != "level" and key != "extracurricularActivities":
                    if key == "university":
                        # Add the value of "university" with bold formatting and without a bullet
                        p = doc.add_paragraph()
                        run = p.add_run("University: " + value)
                        run.bold = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        doc.add_paragraph(f"{value}", style="List Bullet")

    # Title
    add_heading_with_format(doc, "Resume", font_size=18, is_bold=True)

    # Contact Information
    add_heading_with_format(doc, "Contact Information", font_size=16, is_bold=True)
    doc.add_paragraph("Name: " + data["name"])
    doc.add_paragraph("Address: " + data["address"])
    doc.add_paragraph("Email: " + data["email"])
    doc.add_paragraph("LinkedIn: " + data["linkedin"])
    doc.add_paragraph("Phone: " + data["phone"])

    # Education section
    add_details_section(doc, "Education", data["education"])

    # Skills section
    skills = data["skills"]
    skills_text = ", ".join(skill["skills"] for skill in skills)
    add_heading_with_format(doc, "Skills", font_size=14, is_bold=True)
    doc.add_paragraph(skills_text, style="List Bullet")

    # Work Experience section
    add_heading_with_format(doc, "Work Experience", font_size=16, is_bold=True)
    for entry in data["workExperience"]:
        add_details_section(doc, "", [entry], is_bold_title=False)  # Removed the "Work Entry" heading

    # Projects section
    add_heading_with_format(doc, "Projects", font_size=16, is_bold=True)
    for project in data["projects"]:
        add_details_section(doc, "", [project], is_bold_title=False)  # Removed repeated "Project" heading

    # Save the document to a .docx file

    word_buffer = BytesIO()
    output_file_path = "generated_resume.docx"
    doc.save(word_buffer)
    word_buffer.seek(0)

    return word_buffer


@app.route('/resumebuilder', methods=['POST'])
def form_builder():
    try:
        # Assuming the request data is in JSON format
        data = request.json

        # Log the data (you can customize this part)
        print("Received Form Data:")
        for key, value in data.items():
            print(f"{key}: {value}")

        # Generate PDF
        pdf_data = generate_pdf(data)

        # Send the PDF file as a response
        return send_file(pdf_data, mimetype='application/msword', as_attachment=True,
                         attachment_filename='generated_resume.docx')
    except Exception as e:
        print(f"Error processing form data: {str(e)}")
        return "Error processing form data", 500


def get_job_app_status(applications):
    """
    Returns the number of job applications in each job status

    :param: applications: a list of all the job applications the user created
    
    :return: a list containing objects with name and count as the fields for the different statuses the applications are in and the number of applications in those statuses.
    """
    job_app_status = {
        "Wishlist": 0,
        "Applied": 0,
        "Interviewed": 0,
        "Rejected": 0,
    }
    for application in applications:
        # if application["stage"]=="Wishlist":
        if int(application["stage"]) == 1:
            job_app_status["Wishlist"] += 1
        # elif application["stage"]=="Applied":
        elif int(application["stage"]) == 2:
            job_app_status["Applied"] += 1
            job_app_status["Wishlist"] += 1
        # elif application["stage"]=="Interviewed":
        elif int(application["stage"]) == 3:
            job_app_status["Interviewed"] += 1
            job_app_status["Applied"] += 1
            job_app_status["Wishlist"] += 1
        # elif application["stage"]=="Rejected":
        elif int(application["stage"]) == 4:
            job_app_status["Rejected"] += 1
            job_app_status["Interviewed"] += 1
            job_app_status["Applied"] += 1
            job_app_status["Wishlist"] += 1
    res = [
        {"name": "Jobs Saved", "count": job_app_status["Wishlist"]},
        {"name": "Applied", "count": job_app_status["Applied"]},
        {"name": "Interviewed", "count": job_app_status["Interviewed"]},
        {"name": "Rejected", "count": job_app_status["Rejected"]},
    ]
    return res,job_app_status["Applied"],job_app_status["Interviewed"]

def get_last_six_months_job_counts(applications):
    """
    Returns the number of jobs created in each month for the last 6 months

    :param: applications: a list of all the job applications the user created
    
    :return: a list containing objects with Month, Jobs Created as the fields for the last six months
    """
    month_map = {
        index + 1: val
        for index, val in enumerate(
            [
                "Jan", "Feb", "Mar", "Apr",
                "May", "Jun", "Jul", "Aug",
                "Sep", "Oct", "Nov", "Dec",
            ]
        )
    }
    # Create a defaultdict to store date strings for each month
    result_dict = defaultdict(list)
    # Get the current date
    current_date = (datetime.now().replace(day=1) + timedelta(days=32)).replace(day=1)
    # Iterate over the last six months
    for i in range(6):
        # Calculate the start and end date of the current month
        end_of_month = current_date.replace(day=1) - timedelta(days=1)
        start_of_month = end_of_month.replace(day=1)
        # Filter date objects that belong to the current month
        current_month_dates = [
            datetime.strptime(application["date"], "%Y-%m-%d").strftime("%Y-%m-%d")
            for application in applications
            if start_of_month
            <= datetime.strptime(application["date"], "%Y-%m-%d")
            <= end_of_month
        ]
        # Store the result in the dictionary
        result_dict[
            "%s %s" % (month_map[start_of_month.month], start_of_month.year)
        ] = len(current_month_dates)
        # Move to the previous month
        current_date = start_of_month
    res = [
        {"Month": key, "Jobs Created": result_dict[key]}
        for key in list(result_dict.keys())[::-1]
    ]
    return res

def get_last_four_jobs(applications):
    """
    Returns the latest 4 job application details based on the date field in the job application.

    :param: applications: a list of all the job applications the user created
    
    :return: a list containing objects with jobTitle, company and status for the latest 4 job applications based on the date field of the ob application
    """
    appStatus = {'1': "Wishlist", '2': "Applied", '3': "Interviewed", '4':"Rejected"}
    apps = sorted(applications, key=lambda application: datetime.strptime(application["date"], "%Y-%m-%d"), reverse=True)[:4]
    res = [
        {
            "jobTitle": application["jobTitle"],
            "company": application["companyName"],
            "status": appStatus[application["stage"]],
        }
        for application in apps
    ]
    return res

if __name__ == '__main__':
    app.run(debug=False)

# if __name__ == "__main__":
#     app.run()